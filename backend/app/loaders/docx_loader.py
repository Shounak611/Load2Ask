from pathlib import Path
from typing import Union, Dict, Any
from app.loaders.base import BaseLoader
from app.models.internal import Document
from app.core.errors import InvalidFileError
from app.core.logging import logger

try:
    import docx
except ImportError:
    docx = None


class DOCXLoader(BaseLoader):
    """DOCX loader extracting headings, body paragraphs, and structured table data."""

    def load(self, source: Union[str, Path], metadata: Dict[str, Any] = None) -> Document:
        file_path = Path(source)
        if not file_path.exists():
            raise InvalidFileError(f"DOCX file not found: {source}")

        if docx is None:
            raise InvalidFileError("python-docx package is not installed.")

        try:
            doc_obj = docx.Document(str(file_path))
            elements = []
            current_heading = "Introduction"

            # Parse paragraphs and headings
            for p in doc_obj.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                if p.style.name.startswith("Heading"):
                    current_heading = text
                    elements.append(f"\n## {text}\n")
                else:
                    elements.append(text)

            # Parse tables
            for table_idx, table in enumerate(doc_obj.tables):
                table_lines = [f"\n[Table {table_idx + 1}]"]
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):
                        table_lines.append(" | ".join(row_cells))
                elements.append("\n".join(table_lines))

            full_content = "\n".join(elements)

            meta = {
                "source_type": "docx",
                "file_name": file_path.name,
                "paragraph_count": len(doc_obj.paragraphs),
                "table_count": len(doc_obj.tables),
                "file_size": file_path.stat().st_size,
                **(metadata or {})
            }

            return Document(
                source_type="docx",
                source_name=file_path.name,
                source_uri=str(file_path.absolute()),
                content=full_content,
                metadata=meta
            )
        except Exception as e:
            logger.error(f"Failed to load DOCX file {source}: {e}")
            raise InvalidFileError(f"Failed to read DOCX file {source}: {str(e)}")
